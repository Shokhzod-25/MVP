from docx import Document

from app.schemas.schemas import UserData

doc = Document()

async def create_doc_file(user_data: UserData):
    doc.add_paragraph(f"ФИО пользователя: {user_data.username}")
    doc.add_paragraph(f"Возраст: {user_data.age}")
    doc.add_paragraph("Интересы / Комментарии:")

    doc.save(f"{user_data.file_name}.docx")
